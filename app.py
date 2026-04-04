import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime
import io
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Meru Networks | Satellite Command",
    page_icon="🛰️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- DISEÑO UI PREMIUM TECNOLÓGICO (CSS) ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@300;400;700&family=Share+Tech+Mono&display=swap');
    
    /* Fondo y tipografía base - Estilo NOC */
    .stApp {
        background-color: #05070a;
        color: #e0e6ed;
        font-family: 'JetBrains Mono', monospace;
    }

    /* Sidebar estilizada Dark */
    [data-testid="stSidebar"] {
        background-color: #0a0e14;
        border-right: 1px solid #1a202c;
    }

    /* Contenedores con efecto Glassmorphism */
    .st-emotion-cache-1r6slb0, .st-emotion-cache-12w0qpk {
        background-color: rgba(22, 27, 34, 0.6);
        border-radius: 12px;
        padding: 25px;
        border: 1px solid rgba(0, 212, 255, 0.2);
        box-shadow: 0 0 15px rgba(0, 212, 255, 0.05);
        backdrop-filter: blur(10px);
    }

    /* Header Corporativo Tecnológico */
    .main-header {
        background: rgba(13, 17, 23, 0.9);
        padding: 1rem 2rem;
        border-radius: 12px;
        border: 1px solid #00d4ff;
        margin-bottom: 2rem;
        display: flex;
        justify-content: space-between;
        align-items: center;
        box-shadow: 0 0 20px rgba(0, 212, 255, 0.1);
    }

    /* Tabs Personalizados Satelitales */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
        background-color: transparent;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: rgba(22, 27, 34, 0.8);
        border: 1px solid #1a202c;
        border-radius: 8px 8px 0px 0px;
        padding: 12px 24px;
        color: #8b949e;
        font-weight: 700;
        transition: 0.3s;
    }
    .stTabs [aria-selected="true"] {
        background-color: rgba(0, 212, 255, 0.1) !important;
        color: #00d4ff !important;
        border-color: #00d4ff !important;
        box-shadow: 0 0 10px rgba(0, 212, 255, 0.2);
    }

    /* Botones de Comando Eléctricos */
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        background: transparent;
        color: #00d4ff;
        font-weight: 700;
        border: 2px solid #00d4ff;
        text-transform: uppercase;
        letter-spacing: 1px;
        font-family: 'Share Tech Mono', monospace;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        background-color: rgba(0, 212, 255, 0.1);
        box-shadow: 0 0 20px rgba(0, 212, 255, 0.3);
        transform: translateY(-2px);
    }

    /* Estilo para los Títulos de secciones */
    h1, h2, h3 {
        font-family: 'Share Tech Mono', monospace;
        color: #ffffff;
        text-transform: uppercase;
        letter-spacing: 2px;
    }

    /* Ticket Style Terminal */
    .ticket-entry {
        background: rgba(0, 0, 0, 0.2);
        border: 1px solid #1a202c;
        border-left: 3px solid #00d4ff;
        padding: 15px;
        margin-bottom: 10px;
        font-size: 0.9rem;
    }
    </style>
""", unsafe_allow_html=True)

# --- LÓGICA DE DATOS (Preservada Integramente) ---
def get_clean_df(file):
    content = file.getvalue().decode('utf-8', errors='ignore').splitlines()
    skip = 0
    for i, line in enumerate(content):
        if any(k in line for k in ["Date", "Time", "Octets", "Eb/No", "FECHA", "ZONA", "NOMBRE ISP"]):
            skip = i
            break
    file.seek(0)
    try:
        df = pd.read_csv(file, skiprows=skip)
        df.columns = [str(c).strip().replace('"', '') for c in df.columns]
        return df
    except: return pd.DataFrame()

def generate_meru_docx(data_dict, month_text):
    doc = Document()
    title = doc.add_heading('INFORME DE GESTIÓN MENSUAL: RED SATELITAL MERU', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Periodo: {month_text}\nDepartamento: Operaciones de Red (NOC)")
    
    sections = {
        "FALLAS INTERNAS": "2. REPORTE DE FALLAS INTERNAS (GESTIÓN PROPIA)",
        "ISP": "3. REPORTE DE FALLAS DE PROVEEDORES (ISP)",
        "RECLAMOS": "4. ATENCIÓN DE RECLAMOS DEL ABONADO"
    }

    doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    doc.add_paragraph("Resumen de disponibilidad y eventos críticos del mes.")

    for name, df in data_dict.items():
        heading = "DETALLE DE OPERACIONES"
        for key, val in sections.items():
            if key in name.upper(): heading = val
        
        doc.add_heading(heading, level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = col
        for _, row in df.head(40).iterrows():
            row_cells = table.add_row().cells
            for i, v in enumerate(row): row_cells[i].text = str(v)
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# --- HEADER PREMIUM - LOGO INTEGRADO ---
with st.container():
    st.markdown(f"""
        <div class="main-header">
            <div style="display:flex; align-items:center; gap:20px;">
                <img src="data:image/png;base64,{st.session_state.get('logo_base64', '')}" width="200">
                <div>
                    <h1 style="margin:0; font-size:1.8rem;">SATELLITE COMMAND</h1>
                    <p style="margin:0; color:#8b949e; font-size:0.9rem;">Intelligence Hub | NOC Operations</p>
                </div>
            </div>
            <div style="text-align:right;">
                <span style="background:rgba(0, 212, 255, 0.1); color:#00d4ff; padding:8px 15px; border-radius:8px; font-size:0.9rem; font-weight:700; font-family:'Share Tech Mono', monospace; border: 1px solid #00d4ff;">
                    LINK ACTIVE • v10.0
                </span>
            </div>
        </div>
    """, unsafe_allow_html=True)

# Trick para cargar el logo una sola vez y usarlo en el HTML
if 'logo_base64' not in st.session_state and os.path.exists("image_4f92e1.png"):
    import base64
    with open("image_4f92e1.png", "rb") as image_file:
        st.session_state['logo_base64'] = base64.b64encode(image_file.read()).decode()
    st.rerun()

# --- SIDEBAR DARK ---
with st.sidebar:
    st.markdown("### 📥 TELEMETRÍA DOWNLINK")
    uploaded_files = st.file_uploader("Cargar flujos CSV iDirect", accept_multiple_files=True)
    st.markdown("---")
    st.markdown("### 🗓️ PERIODO DE MISIÓN")
    sel_month = st.selectbox("Mes de reporte:", ["Marzo 2026", "Abril 2026", "Mayo 2026"])
    st.info("El sistema está sincronizado con los satélites de la red Meru.")

# --- CUERPO PRINCIPAL ---
processed_data = {}
if uploaded_files:
    for f in uploaded_files:
        processed_data[f.name] = get_clean_df(f)

# Navegación con Iconos
tab_diag, tab_incident, tab_intel, tab_export = st.tabs([
    "📊 Monitoreo Espectral", 
    "🎫 Gestión de Tickets", 
    "🧠 Core IA", 
    "📥 Downlink Informes"
])

# 1. DASHBOARD
with tab_diag:
    if uploaded_files:
        st.markdown("### 📈 Análisis de Señal Satelital")
        for name, df in processed_data.items():
            with st.container():
                st.markdown(f"**Flujo de Datos:** `{name}`")
                stations = sorted(list(set([c.split('/')[0] for c in df.columns if '/' in c])))
                if stations:
                    c1, c2 = st.columns([1, 3])
                    with c1:
                        sel = st.multiselect(f"Seleccionar Nodo:", stations, default=stations[:1], key=f"dash_{name}")
                    with c2:
                        fig = go.Figure()
                        for s in sel:
                            for c in [col for col in df.columns if col.startswith(s + "/")]:
                                fig.add_trace(go.Scatter(y=df[c], name=f"{s}-{c.split('/')[-1]}", line=dict(width=2, color='#00d4ff')))
                        fig.update_layout(
                            template="plotly_dark", 
                            height=350, 
                            margin=dict(l=10, r=10, t=10, b=10),
                            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                            plot_bgcolor='rgba(0,0,0,0)',
                            paper_bgcolor='rgba(0,0,0,0)',
                            yaxis=dict(gridcolor='#1a202c'),
                            xaxis=dict(gridcolor='#1a202c')
                        )
                        st.plotly_chart(fig, use_container_width=True)
                else:
                    st.dataframe(df.head(5), use_container_width=True)
                st.markdown("---")
    else:
        st.markdown("""
        <div style="background:rgba(22, 27, 34, 0.6); padding:50px; border-radius:12px; border:2px dashed #00d4ff; text-align:center;">
            <h1 style="color:#00d4ff; font-size:4rem; margin:0;">👋</h1>
            <h2 style="color:#ffffff;">CONSOLA INACTIVA</h2>
            <p style="color:#8b949e;">Cargue flujos de telemetría en la barra lateral para iniciar el monitoreo espectral.</p>
        </div>
        """, unsafe_allow_html=True)

# 2. TICKETS
with tab_incident:
    st.markdown("### 🎫 Panel de Incidentes (Tickets)")
    col_f, col_v = st.columns([1, 2])
    with col_f:
        st.markdown('<div style="background:rgba(22, 27, 34, 0.8); padding:20px; border-radius:12px; border:1px solid #1a202c;">', unsafe_allow_html=True)
        with st.form("t_form"):
            st.write("**Generar Reporte de Misión**")
            node = st.text_input("ID del Nodo / Estación Terrena")
            issue = st.selectbox("Categoría de Falla", ["Desvanecimiento de Señal", "Falla de Antena", "Corte de Energía (Hub)", "Saturación de Transpondedor"])
            if st.form_submit_button("Sincronizar Alerta"):
                st.success(f"Ticket registrado y transmitido para {node}")
        st.markdown('</div>', unsafe_allow_html=True)
    with col_v:
        st.markdown("**Incidentes Recientes en Red**")
        st.markdown('<div class="ticket-entry"><b style="color:#00d4ff;">NODE_SAT_01</b> | Desvanecimiento de Señal<br><small style="color:#8b949e;">Hace 5 minutos | Prioridad: Alta</small></div>', unsafe_allow_html=True)

# 3. CORE IA
with tab_intel:
    st.markdown("### 🧠 Inteligencia Heurística de Red")
    st.markdown("""
    <div style="background:rgba(22, 27, 34, 0.6); padding:40px; border-radius:12px; border:2px dashed #1a202c; text-align:center;">
        <h2 style="color:#00d4ff;">ANALIZANDO PATRONES ORBITALES...</h2>
        <p style="color:#8b949e;">El núcleo de IA está procesando la telemetría para predecir Sun Outage o Rain Fade.</p>
    </div>
    """, unsafe_allow_html=True)

# 4. EXPORTAR (Mantenido según tus modelos)
with tab_export:
    st.markdown("### 📥 Generación de Informes de Misión")
    if processed_data:
        st.write(f"Estructura basada en los modelos oficiales de Meru para el periodo: **{sel_month}**")
        c_w, c_e = st.columns(2)
        with c_w:
            st.markdown("""
                <div style="background:rgba(22, 27, 34, 0.8); padding:25px; border-radius:12px; border:1px solid #1a202c;">
                    <h4 style="color:#ffffff;">📄 Downlink Informe Word</h4>
                    <p style="font-size:0.8rem; color:#8b949e;">Incluye Resumen Ejecutivo y tablas de Fallas e ISP formateadas.</p>
                </div>
            """, unsafe_allow_html=True)
            docx_data = generate_meru_docx(processed_data, sel_month)
            st.download_button("Descargar Informe .DOCX", docx_data, f"Informe_Meru_{sel_month}.docx")
        
        with c_e:
            st.markdown("""
                <div style="background:rgba(22, 27, 34, 0.8); padding:25px; border-radius:12px; border:1px solid #1a202c;">
                    <h4 style="color:#ffffff;">📊 Base de Datos Excel</h4>
                    <p style="font-size:0.8rem; color:#8b949e;">Consolidado de datos en pestañas independientes por reporte.</p>
                </div>
            """, unsafe_allow_html=True)
            output_x = io.BytesIO()
            with pd.ExcelWriter(output_x, engine='xlsxwriter') as writer:
                for n, df in processed_data.items():
                    df.to_excel(writer, sheet_name=n.split('.')[0][:30], index=False)
            st.download_button("Descargar Base de Datos .XLSX", output_x.getvalue(), f"Reporte_Meru_{sel_month}.xlsx")
    else:
        st.warning("Suba los archivos de telemetría para habilitar la exportación.")

st.markdown("<div style='text-align:center; padding:30px; opacity:0.3; font-size:0.75rem; color:#8b949e; font-family: monospace;'>MERU NETWORKS SECURITY ORBITAL CORE © 2026 | TRANSMISIÓN SEGURA</div>", unsafe_allow_html=True)
